"""Convert a subset of Markdown to Word (.docx) for human readers."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_formatted_run(paragraph, text: str, bold: bool = False, code: bool = False) -> None:
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(part)
            if bold:
                run.bold = True
            if code:
                run.font.name = 'Consolas'


def _parse_table_rows(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        if not line.strip().startswith('|'):
            break
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        if all(re.fullmatch(r'-+', c.replace(' ', '')) for c in cells):
            continue
        rows.append(cells)
    return rows


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding='utf-8')
    lines = text.splitlines()
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith('```'):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.25)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped == '---':
            doc.add_paragraph('')
            i += 1
            continue

        if stripped.startswith('# '):
            doc.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue
        if stripped.startswith('## '):
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:].strip(), level=3)
            i += 1
            continue

        if stripped.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            rows = _parse_table_rows(table_lines)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Table Grid'
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        _add_formatted_run(p, cell_text, bold=(r_idx == 0))
                        if r_idx == 0:
                            _set_cell_shading(cell, 'E8EEF4')
                doc.add_paragraph('')
            continue

        if stripped.startswith('- '):
            while i < len(lines) and lines[i].strip().startswith('- '):
                item = lines[i].strip()[2:]
                p = doc.add_paragraph(style='List Bullet')
                _add_formatted_run(p, item)
                i += 1
            continue

        if re.match(r'^\d+\.\s', stripped):
            while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
                item = re.sub(r'^\d+\.\s', '', lines[i].strip())
                p = doc.add_paragraph(style='List Number')
                _add_formatted_run(p, item)
                i += 1
            continue

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        _add_formatted_run(p, stripped)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f'Wrote {docx_path}')


def main(argv: list[str] | None = None) -> int:
    argv = argv or sys.argv[1:]
    root = Path(__file__).resolve().parents[1]
    targets = argv or [
        'docs/ARCHITECTURE.md',
        'docs/POTENTIAL_FUTURE_FUNCTIONALITY.md',
    ]
    for rel in targets:
        md_path = root / rel
        docx_path = md_path.with_suffix('.docx')
        if not md_path.exists():
            print(f'Missing: {md_path}', file=sys.stderr)
            return 1
        md_to_docx(md_path, docx_path)
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
